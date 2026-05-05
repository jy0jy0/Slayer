"""Interview Prep page — generate tailored interview questions by category."""

import datetime
import json
import logging
import time
from io import BytesIO

import streamlit as st

logger = logging.getLogger(__name__)
from slayer.ui.styles import GLOBAL_CSS
from slayer.ui.components import render_page_header


# Category display config: (label, emoji)
_CATEGORY_DISPLAY = {
    "기술": ("Technical", "💻"),
    "경험": ("Experience", "📋"),
    "상황/행동": ("Situational / Behavioral", "🎭"),
    "인성": ("Personality", "🧠"),
    "컬처핏": ("Culture Fit", "🤝"),
    "기업 이해도": ("Company Knowledge", "🏢"),
}

_REQUIRES_RESEARCH = {"컬처핏", "기업 이해도"}


def _build_docx(result, jd) -> bytes:
    """InterviewQuestionsOutput을 DOCX 바이트로 변환."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(f"{jd.company} — {jd.position} 면접 준비 자료", 0)
    doc.add_paragraph(f"생성일: {datetime.date.today()}")

    if result.weak_areas:
        doc.add_heading("우선 대비 영역", level=1)
        for area in result.weak_areas:
            doc.add_paragraph(area, style="List Bullet")

    by_category: dict[str, list] = {}
    for q in result.questions:
        cat = q.category.value if hasattr(q.category, "value") else str(q.category)
        by_category.setdefault(cat, []).append(q)

    for cat, questions in by_category.items():
        display_label, _ = _CATEGORY_DISPLAY.get(cat, (cat, ""))
        doc.add_heading(f"{display_label} ({cat})", level=1)
        for i, q in enumerate(questions, 1):
            p = doc.add_paragraph()
            p.add_run(f"Q{i}. {q.question}").bold = True
            doc.add_paragraph(f"💡 Tip: {q.tip}")
            doc.add_paragraph()

    if result.sample_answers:
        doc.add_heading("예시 답변", level=1)
        for sa in result.sample_answers:
            p = doc.add_paragraph()
            p.add_run(f"Q. {sa.question}").bold = True
            doc.add_paragraph(f"A. {sa.answer}")
            doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _run_generation(jd, resume, cr, mr, questions_per_category: int) -> None:
    """최초 질문 생성."""
    t_start = time.time()
    with st.status("Generating interview questions...", expanded=True) as status:
        try:
            from slayer.pipelines.interview_questions import generate_interview_questions
            from slayer.schemas import InterviewQuestionsInput

            status.write("⏳ Preparing input data...")
            inp = InterviewQuestionsInput(
                jd=jd,
                resume=resume,
                company_research=cr,
                match_result=mr,
                questions_per_category=questions_per_category,
            )

            status.write("⏳ Calling LLM to generate questions...")
            result = generate_interview_questions(inp)

            if "interview_result" in st.session_state:
                st.session_state["interview_history"].append(
                    st.session_state["interview_result"]
                )
            st.session_state["interview_result"] = result

            status.write(f"✅ Generated **{len(result.questions)}** questions across categories")
            if result.excluded_categories:
                status.write(
                    f"ℹ️ Excluded categories (insufficient data): {', '.join(result.excluded_categories)}"
                )
            status.update(label="✅ Interview questions generated", state="complete")

            duration_ms = int((time.time() - t_start) * 1000)
            try:
                from slayer.db.repository import save_agent_log, save_interview_questions
                save_interview_questions(jd, result)
                save_agent_log(
                    agent_name="interview_prep",
                    status="success",
                    input_summary=f"questions_per_category={questions_per_category}, company={jd.company}, title={jd.title}",
                    output_summary=f"total_questions={len(result.questions)}, excluded={result.excluded_categories or []}, weak_areas={len(result.weak_areas) if result.weak_areas else 0}",
                    duration_ms=duration_ms,
                )
            except Exception as e:
                logger.warning("DB save failed: %s", e)

        except Exception as e:
            status.update(label="❌ Generation failed", state="error")
            st.error(f"Generation failed: {e}")
            try:
                from slayer.db.repository import save_agent_log
                save_agent_log(agent_name="interview_prep", status="failed", error_message=str(e)[:500])
            except Exception:
                pass
            return

    st.rerun()


def _run_refinement(current_result, feedback_text: str, focus_cat_values: list[str], jd, resume, cr, mr, questions_per_category: int) -> None:
    """피드백 기반 재생성."""
    with st.status("Refining questions...", expanded=True) as status:
        try:
            from slayer.pipelines.interview_questions import refine_interview_questions
            from slayer.schemas import InterviewCategory, InterviewQuestionsInput, RefinementFeedback

            pinned_qs = [
                q for q in current_result.questions
                if st.session_state["interview_pinned"].get(q.question, False)
            ]

            inp = InterviewQuestionsInput(
                jd=jd,
                resume=resume,
                company_research=cr,
                match_result=mr,
                questions_per_category=questions_per_category,
            )

            focus_categories = [InterviewCategory(v) for v in focus_cat_values]
            feedback = RefinementFeedback(
                free_text=feedback_text,
                focus_categories=focus_categories,
                pinned_questions=pinned_qs,
            )

            status.write(f"⏳ Refining with {len(pinned_qs)} pinned questions...")
            new_result = refine_interview_questions(inp, feedback)

            st.session_state["interview_history"].append(current_result)
            st.session_state["interview_result"] = new_result

            status.write(f"✅ Refined — {len(new_result.questions)} questions (pinned {len(pinned_qs)} + new {len(new_result.questions) - len(pinned_qs)})")
            status.update(label="✅ Refinement complete", state="complete")

            try:
                from slayer.db.repository import save_agent_log, save_interview_questions
                save_interview_questions(jd, new_result)
                save_agent_log(
                    agent_name="interview_prep_refine",
                    status="success",
                    input_summary=f"iteration={len(st.session_state['interview_history'])}, pinned={len(pinned_qs)}, feedback_len={len(feedback_text)}",
                    output_summary=f"total_questions={len(new_result.questions)}",
                )
            except Exception as e:
                logger.warning("DB save failed: %s", e)

        except Exception as e:
            status.update(label="❌ Refinement failed", state="error")
            st.error(f"Refinement failed: {e}")
            return

    st.rerun()


def render():
    st.html(GLOBAL_CSS)
    render_page_header("Interview Prep", "Generate tailored interview questions based on JD and resume.")

    # ── Session state init ────────────────────────────────────────────
    if "interview_pinned" not in st.session_state:
        st.session_state["interview_pinned"] = {}
    if "interview_notes" not in st.session_state:
        st.session_state["interview_notes"] = {}
    if "interview_history" not in st.session_state:
        st.session_state["interview_history"] = []

    # ── Prerequisites ─────────────────────────────────────────────────
    has_jd = "jd_data" in st.session_state
    has_resume = "resume_data" in st.session_state
    has_research = "company_research" in st.session_state
    has_match = "match_result" in st.session_state

    cols = st.columns(4)
    with cols[0]:
        st.success("JD loaded") if has_jd else st.warning("JD not loaded")
    with cols[1]:
        st.success("Resume loaded") if has_resume else st.warning("Resume not loaded")
    with cols[2]:
        if has_research:
            st.success(f"Research: {st.session_state['company_research'].company_name}")
        else:
            st.info("No research (optional)")
    with cols[3]:
        if has_match:
            st.success(f"Match: ATS {st.session_state['match_result'].ats_score:.0f}")
        else:
            st.info("No match (optional)")

    if not has_jd or not has_resume:
        st.warning("JD and Resume are required. Please run JD-Resume Match first or provide data.")
        return

    # ── Config ────────────────────────────────────────────────────────
    questions_per_category = st.slider(
        "Questions per category", min_value=1, max_value=10, value=3, key="interview_qpc"
    )

    # ── Generate + Reset 버튼 ─────────────────────────────────────────
    col_gen, col_reset = st.columns([3, 1])
    with col_gen:
        run_btn = st.button("🎯 Generate Interview Questions", type="primary", use_container_width=True)
    with col_reset:
        reset_btn = st.button("🗑 Reset", use_container_width=True)

    if reset_btn:
        for key in ["interview_result", "interview_history", "interview_pinned", "interview_notes"]:
            st.session_state.pop(key, None)
        st.rerun()

    if run_btn:
        from slayer.schemas import JDSchema, ParsedResume
        jd = JDSchema(**json.loads(st.session_state["jd_data"]))
        resume = ParsedResume(**json.loads(st.session_state["resume_data"]))
        cr = st.session_state.get("company_research")
        mr = st.session_state.get("match_result")
        _run_generation(jd, resume, cr, mr, questions_per_category)

    if "interview_result" not in st.session_state:
        st.info("Click the button above to generate interview questions.")
        return

    st.divider()
    result = st.session_state["interview_result"]

    # JD 객체 (배너/다운로드/refine에서 재사용)
    from slayer.schemas import JDSchema, ParsedResume
    jd = JDSchema(**json.loads(st.session_state["jd_data"]))
    resume = ParsedResume(**json.loads(st.session_state["resume_data"]))
    cr = st.session_state.get("company_research")
    mr = st.session_state.get("match_result")

    # ── Excluded categories notice ────────────────────────────────────
    if result.excluded_categories:
        st.warning(
            f"Excluded categories due to missing data: **{', '.join(result.excluded_categories)}**. "
            "Provide Company Research and Match Result to enable all categories."
        )

    # ── 이터레이션 배너 + DOCX 다운로드 ──────────────────────────────
    iteration_n = len(st.session_state["interview_history"]) + 1
    pinned_count = sum(
        1 for q in result.questions
        if st.session_state["interview_pinned"].get(q.question, False)
    )

    col_banner, col_dl = st.columns([4, 1])
    with col_banner:
        st.html(
            f'<div style="display:flex; gap:24px; padding:12px 16px; background:#1a2744; '
            f'border-radius:8px; margin-bottom:4px; align-items:center;">'
            f'<span style="color:#93b5e8; font-size:13px;">Iteration '
            f'<strong style="color:#fff;">{iteration_n}</strong></span>'
            f'<span style="color:#93b5e8; font-size:13px;">Questions '
            f'<strong style="color:#fff;">{len(result.questions)}</strong></span>'
            f'<span style="color:#93b5e8; font-size:13px;">Pinned '
            f'<strong style="color:#fff;">{pinned_count}</strong></span>'
            f'</div>'
        )
    with col_dl:
        docx_bytes = _build_docx(result, jd)
        filename = f"{jd.company}_{jd.position}_면접질문.docx".replace(" ", "_")
        st.download_button(
            label="⬇ DOCX",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    # ── Questions by category ─────────────────────────────────────────
    st.markdown("### Questions by Category")

    by_category: dict[str, list] = {}
    for q in result.questions:
        cat = q.category.value if hasattr(q.category, "value") else str(q.category)
        by_category.setdefault(cat, []).append(q)

    for cat, questions in by_category.items():
        display_label, emoji = _CATEGORY_DISPLAY.get(cat, (cat, "❓"))
        with st.expander(f"{emoji} {display_label} ({cat}) — {len(questions)} questions", expanded=True):
            for i, q in enumerate(questions, 1):
                q_key = f"{cat}_{i}"
                pin_state = st.session_state["interview_pinned"].get(q.question, False)

                col_q, col_pin = st.columns([11, 1])
                with col_q:
                    pin_icon = " 📌" if pin_state else ""
                    st.markdown(f"**Q{i}. {q.question}**{pin_icon}")
                with col_pin:
                    btn_label = "✅" if pin_state else "📌"
                    if st.button(btn_label, key=f"pin_{q_key}", help="Pin to keep across refinements"):
                        st.session_state["interview_pinned"][q.question] = not pin_state
                        st.rerun()

                # Tip — 항상 표시
                st.html(
                    f'<div style="background:#1a2744; border-left:3px solid #3b82f6; '
                    f'padding:8px 12px; margin:4px 0 8px 0; border-radius:4px; '
                    f'font-size:13px; color:#93b5e8;">💡 Tip: {q.tip}</div>'
                )

                # Intent — 기본 접힘
                if st.toggle("Show intent", key=f"intent_{q_key}", value=False):
                    st.caption(f"Intent: {q.intent}")

                # Source badge
                st.html(
                    f'<span style="display:inline-block; background:#2a2a3e; color:#888; '
                    f'padding:2px 8px; border-radius:4px; font-size:11px; margin-bottom:8px;">'
                    f'Source: {q.source}</span>'
                )

                # 개선 방향 메모
                existing_note = st.session_state["interview_notes"].get(q.question, "")
                note = st.text_input(
                    "개선 방향",
                    value=existing_note,
                    key=f"note_{q_key}",
                    placeholder="e.g. 더 구체적인 기술 질문으로, STAR 구조 강조 등",
                    label_visibility="collapsed",
                )
                if note != existing_note:
                    st.session_state["interview_notes"][q.question] = note

    # ── Refine Questions 폼 ───────────────────────────────────────────
    st.divider()
    st.markdown("### Refine Questions")

    # focus_categories 중 데이터 부족으로 제외되는 카테고리 경고
    if not has_research:
        st.caption(
            "ℹ️ 컬처핏/기업이해도 카테고리는 Company Research 없이 선택해도 자동 제외됩니다."
        )

    with st.form("feedback_form"):
        feedback_text = st.text_area(
            "Feedback",
            placeholder="e.g. 기술 질문을 더 심화해줘. Kubernetes 관련 질문 추가. 인성 질문은 덜 일반적으로.",
            height=100,
            label_visibility="collapsed",
        )
        focus_cats = st.multiselect(
            "Focus categories (leave empty for all)",
            options=list(_CATEGORY_DISPLAY.keys()),
            default=[],
            key="feedback_focus_cats",
        )

        pinned_qs = [
            q for q in result.questions
            if st.session_state["interview_pinned"].get(q.question, False)
        ]
        if pinned_qs:
            st.caption(f"📌 {len(pinned_qs)} question(s) pinned — will be preserved")

        refine_btn = st.form_submit_button("🔄 Refine Questions", type="primary", use_container_width=True)

    if refine_btn:
        if not feedback_text.strip() and not pinned_qs and not focus_cats:
            st.warning("Please provide feedback, pin questions, or select focus categories.")
        else:
            _run_refinement(result, feedback_text, focus_cats, jd, resume, cr, mr, questions_per_category)

    # ── Sample Answers ────────────────────────────────────────────────
    if result.sample_answers:
        st.divider()
        st.markdown("### Sample Answers")
        for sa in result.sample_answers:
            with st.expander(f"💬 {sa.question}"):
                st.markdown(sa.answer)

    # ── Weak Areas / Priority ─────────────────────────────────────────
    if result.weak_areas:
        st.markdown("### Priority Areas")
        st.caption("Areas to focus your preparation on.")
        for area in result.weak_areas:
            st.markdown(f"- ⚠️ {area}")

    # ── Previous Iterations ───────────────────────────────────────────
    history = st.session_state["interview_history"]
    if history:
        st.divider()
        st.markdown("### Previous Iterations")
        for idx, old_result in enumerate(reversed(history)):
            iter_num = len(history) - idx
            with st.expander(
                f"Iteration {iter_num} — {len(old_result.questions)} questions",
                expanded=False,
            ):
                old_by_cat: dict[str, list] = {}
                for old_q in old_result.questions:
                    old_cat = old_q.category.value if hasattr(old_q.category, "value") else str(old_q.category)
                    old_by_cat.setdefault(old_cat, []).append(old_q)
                for old_cat, old_qs in old_by_cat.items():
                    old_label, old_emoji = _CATEGORY_DISPLAY.get(old_cat, (old_cat, "❓"))
                    st.markdown(f"**{old_emoji} {old_label}**")
                    for old_q in old_qs:
                        st.markdown(f"- {old_q.question}")

    # ── Raw JSON ──────────────────────────────────────────────────────
    with st.expander("📋 Raw JSON"):
        st.json(result.model_dump())
